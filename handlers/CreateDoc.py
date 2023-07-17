import Buttons, Calls, Messages, config, database, os

from telebot import types
from bot import bot, FilesDatabase, Users, drive
from pyDrive import GDrive


questions = {
    '1': 'Дата заполнения',
    '2': 'Название дня',
    '3': 'Цели дня',
    '4': 'План на день',
    '5': 'Анализ дня, выводы, размышления',
    '6': 'Заполнял',
    '7': 'Название документа'
}


def ReadDoc(query: types.CallbackQuery):
    user = query.message.chat.id
    files = os.listdir(config.DiaryCache)
    if str(user) not in files:
        import datetime
        path = config.DiaryCache + '/' + str(user)
        os.mkdir(path)
        for quest in questions.keys():
            filepath = path + '/' + quest + '.txt'
            if quest == '1':
                answer = '.'.join(str(datetime.date.today()).split('-'))
            elif quest == '6':
                answer = Users.Get(str(user))[1]
            elif quest == '7':
                answer = f"Запись от {'.'.join(str(datetime.date.today()).split('-'))}"
            else:
                answer = 'Пусто'
            with open(filepath, 'w+') as file:
                file.write(answer)
        bot.send_message(chat_id=user, text='Создан новый документ!')
    else:
        bot.send_message(chat_id=user, text='Была найдена незаконченная запись:')
    string = 'Ваш ответ на данный момент:\n'
    keyboard = types.InlineKeyboardMarkup(row_width=6)
    row = []
    for quest in range(1, 8):
        path = config.DiaryCache + '/' + str(user) + '/' + str(quest) + '.txt'
        with open(path, 'r') as file:
            text = file.read()
        if len(text) > 25:
            text = text[:25] + '...'
        string += f'\n{quest} ({questions[str(quest)]}): {text}'
        row.append(types.InlineKeyboardButton(text=str(quest), callback_data=f'EDIT{quest}'))
    keyboard.row(*row)
    keyboard.row(
        types.InlineKeyboardButton(text=Buttons.save, callback_data=Calls.DocSave),
        types.InlineKeyboardButton(text=Buttons.question, callback_data=Calls.EditInfo),
        types.InlineKeyboardButton(text=Buttons.erase, callback_data=Calls.EraseDoc)
    )
    keyboard.row(types.InlineKeyboardButton(text=Buttons.main, callback_data=Calls.Menu))
    string += f'''
\n1-6: Редактировать вопрос
{Buttons.save} - Сохранить документ
{Buttons.question} - Получить подсказку
{Buttons.erase} - Удалить файл
'''
    bot.send_message(chat_id=user, text=string, reply_markup=keyboard)
    

def EditQuestion(query: types.CallbackQuery):
    user = query.message.chat.id
    import database 
    Operations = database.DB('Database/Operation.txt')
    Operations.Delete(str(user))
    Operations.AddString(str(user), mass=['EditQuestion'])
    data = query.data[4:]
    global QuestionNow
    QuestionNow = data
    bot.send_message(chat_id=user, text=f'Пришлите ответ на вопрос {questions[data]}:')


def SaveAnswer(message: types.Message):
    user = message.chat.id
    path = config.DiaryCache + '/' + str(user) + '/' + str(QuestionNow) + '.txt'
    with open(path, 'w') as file:
        file.write(str(message.text))
    keyboard = types.InlineKeyboardMarkup(row_width=1)
    keyboard.add(
        types.InlineKeyboardButton(text=Buttons.back, callback_data=Calls.CreateFile),
        types.InlineKeyboardButton(text=Buttons.main, callback_data=Calls.Menu)
    )
    bot.send_message(chat_id=user, text='Ответ отредактирован!', reply_markup=keyboard)


def SaveDoc(query: types.CallbackQuery):
    user = query.message.chat.id

    from docx import Document
    from docx.shared import Pt
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.style import WD_STYLE_TYPE
    document = Document()

    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    for quest in range(1, 7):
        p = document.add_paragraph('')
        p.add_run(questions[str(quest)]).bold = True
        
        path = config.DiaryCache + '/' + str(user) + '/' + str(quest) + '.txt'
        with open(path, 'r') as file:
            text = file.read()
        p = document.add_paragraph('')
        p.add_run(text)

    path = config.DiaryCache + '/' + str(user) + '/' + str(7) + '.txt'
    with open(path, 'r') as file:
            filename = file.read()
    filename = str(filename) + '.docx'
    path = config.DiaryCache + '/' + filename
    document.save(path)
    db = database.DB(path=config.DiaryDatabase)
    FolderID = db.Get(key=str(user))[-2]
    drive.SendFile(file_path=path, title=filename, folder_parent=FolderID)
    methods = database.DB(path=config.MethodsDatabase)
    for method in methods.data.keys():
        bot.send_document(
            chat_id=str(method), 
            document=open(path, 'rb'), 
            caption=f'Получен новый документ от {db.Get(key=str(user))[-1]}: {filename}'
        )

    os.remove(path)
    directory = config.DiaryCache + '/' + str(user)
    for file in os.listdir(directory):
        os.remove(config.DiaryCache + '/' + str(user) + '/' + file)
    os.rmdir(directory)

    keyboard = types.InlineKeyboardMarkup(row_width=1)
    keyboard.add(types.InlineKeyboardButton(text=Buttons.back, callback_data=Calls.ShowDiary))
    bot.send_message(text=f'Документ под названием {filename} успешно сохранен!', chat_id=user, reply_markup=keyboard)


def EraseDoc(query: types.CallbackQuery):
    user = query.message.chat.id
    directory = config.DiaryCache + '/' + str(user)
    for file in os.listdir(directory):
        os.remove(config.DiaryCache + '/' + str(user) + '/' + file)
    os.rmdir(directory)
    keyboard = types.InlineKeyboardMarkup(row_width=1)
    keyboard.add(types.InlineKeyboardButton(text=Buttons.back, callback_data=Calls.Diary))
    bot.send_message(text=f'Успешно!', chat_id=user, reply_markup=keyboard)


def EditInfo(query: types.CallbackQuery):
    user = query.message.chat.id
    keyboard = types.InlineKeyboardMarkup(row_width=1)
    keyboard.add(types.InlineKeyboardButton(text=Buttons.back, callback_data=Calls.CreateFile))
    bot.send_photo(chat_id=user, photo=open('Database/План анализа.jpg', 'rb'), caption='Лови подсказку!', reply_markup=keyboard)


def handler():
    bot.register_callback_query_handler(callback=ReadDoc, func=lambda query: query.data == Calls.CreateFile)
    bot.register_callback_query_handler(callback=EditQuestion, func=lambda query: query.data[:4] == 'EDIT')
    bot.register_callback_query_handler(callback=SaveDoc, func=lambda query: query.data == Calls.DocSave)
    bot.register_callback_query_handler(callback=EraseDoc, func=lambda query: query.data == Calls.EraseDoc)
    bot.register_callback_query_handler(callback=EditInfo, func=lambda query: query.data == Calls.EditInfo)